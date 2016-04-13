import pyodbc
from docx import Document
import io
from PIL import Image
import sys

filename = raw_input("Path to Access Database(eg Database.mdb): ")

con = pyodbc.connect("Driver={Microsoft Access Driver (*.mdb, *.accdb)};Dbq=./"+filename+";Uid=;Pwd=;")

cursor = con.cursor()

output = Document()


cursor.execute("select IDWclass,count(IDWclass) as tally from IDWPPSLEKKICARD group by IDWclass")
classFrequecy = cursor.fetchall()


print "Writing Classes Table"
for classStd in classFrequecy:
  if classStd[0] is None:
    continue

  p = output.add_paragraph()
  p.add_run(classStd[0])
  p.underline = True
  p.add_run("\n")
  p.add_run("Total: " + str(classStd[1]))
  
  cursor.execute("select * from IDWPPSLEKKICARD where IDWclass = '" + str(classStd[0]) + "' order by IDWsurname asc")
  students = cursor.fetchall()

  table = output.add_table(rows=1,cols=5)
  hdr = table.rows[0].cells
  hdr[0].text = "ID No"
  hdr[1].text = "Last Name"
  hdr[2].text = "First Name"
  hdr[3].text = "Class"
  hdr[4].text = "Picture"
  count = 1

  print "Class: " + classStd[0]
  for student in students:
    print count,
    rows_cells = table.add_row().cells
    rows_cells[0].text = student[0]
    rows_cells[1].text = student[2]
    rows_cells[2].text = student[3]
    rows_cells[3].text = student[4] if student[4] is not None else ""
    if student[1] is None:
      continue
    p = rows_cells[4].paragraphs[0]
    pic = Image.open(io.BytesIO(student[1]))
    pic.thumbnail('128','128')
    run = p.add_run()
    pic.save('./pic.jpeg')
    run.add_picture('./pic.jpeg',height=1200000)

    count += 1
  output.add_page_break()
  print 
  print 
output.save('test.docx')

print "Done"
